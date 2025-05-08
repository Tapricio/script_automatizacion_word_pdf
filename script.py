import os
import re
import pandas as pd
from datetime import datetime
from tkinter import Tk, filedialog
from docx import Document
import win32com.client


# Obtener la fecha actual en formato DD-MM-YYYY
fecha_hoy = datetime.today().strftime('%d-%m-%Y')
idTemplate = 10001

# Crear las carpetas necesarias sin sobrescribir
#base_folder_name = rf"G:\Unidades compartidas\Salud\Cartas devolución jubilados\{fecha_hoy} carta devolución jubilados"
base_folder_name = rf"G:\Unidades compartidas\Informática\Cartas devolución jubilados\{fecha_hoy} carta devolución jubilados"
base_folder = base_folder_name
counter = 1
while os.path.exists(base_folder):
    base_folder = f"{base_folder_name} ({counter})"
    counter += 1

pdf_folder = os.path.join(base_folder, f"{fecha_hoy} pdf")
word_folder = os.path.join(base_folder, f"{fecha_hoy} word")

os.makedirs(pdf_folder, exist_ok=True)
os.makedirs(word_folder, exist_ok=True)

# Oculta la ventana principal de Tkinter
root = Tk()
root.withdraw()

# Selección de archivo (descomentar para usar ventana de diálogo)

file_path = filedialog.askopenfilename(
    title="Selecciona un archivo Excel",
    filetypes=[("Excel files", "*.xlsx *.xls")]
)

fechasInvalidas = []
fechasValidas = []
rutInvalidos = []
edadInvalidas = []
casosInvalidos = []
idInvalidos = []
desdentadoSiNombreDocx = "plantillaSinCosto.docx"
desdentadoNoNombreDocx = "plantillaConCosto.docx"

val = 0

def limpiar_nombre_archivo(nombre):
    return re.sub(r'[\/:*?"<>|]', '', nombre)

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)




def docx_replace_multiple_regex(doc_obj, replacements):
    # Compila todos los regex
    compiled_replacements = {re.compile(k): v for k, v in replacements.items()}

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            for regex, replace in compiled_replacements.items():
                if regex.search(p.text):
                    inline = p.runs
                    for i in range(len(inline)):
                        if regex.search(inline[i].text):
                            inline[i].text = regex.sub(replace, inline[i].text)

    def replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

    replace_in_paragraphs(doc_obj.paragraphs)
    replace_in_tables(doc_obj.tables)



if file_path:
    df = pd.read_excel(file_path, header=1)
    df.iloc[:, 4] = pd.to_datetime(df.iloc[:, 4], errors='coerce')  # Columna fecha nacimiento

    for index, row in df.iterrows():
        if pd.notna(df.iloc[index, 0]):  # Validar que haya número de paciente
            fechaNacimiento = df.iloc[index, 4]

            if pd.notna(fechaNacimiento):
                try:
                    nombre = re.sub(r'\s+', ' ', df.iloc[index, 2].strip())

                    rut_response = df.iloc[index, 3].replace(" ", "").replace(".", "")
                    rut_base = rut_response[:-2]
                    verificador = rut_response[-1]
                    rut = "{:,}".format(int(rut_base)).replace(",", ".") + "-" + verificador
                    rut = rut.upper()

                    try:
                        fechaFormateada = fechaNacimiento.strftime('%d-%m-%Y')
                    except Exception:
                        fechasInvalidas.append(index)
                        continue

                    try:
                        edad_raw = df.iloc[index, 5]
                        if isinstance(edad_raw, str):
                            edad_raw = edad_raw.strip()
                        edad = int(float(edad_raw))
                    except Exception:
                        edadInvalidas.append(index)
                        continue

                    print(f"Paciente: {nombre}, Rut: {rut}, Fecha de Nacimiento: {fechaFormateada}, Edad: {str(edad)}, ID: {str(idTemplate)}")
                    val += 1
                    fechasValidas.append(index)
                    
                    try:
                        desdentado = df.iloc[index,1].lower().strip().replace(" ","")
                        if desdentado=="desdentado":
                            #doc = Document(desdentadoSiNombreDocx)
                            doc = Document(rf"G:\Unidades compartidas\Informática\Template cartas devolución jubilados\{desdentadoSiNombreDocx}")
                        
                        else:
                            #doc = Document(desdentadoNoNombreDocx)
                            doc = Document(rf"G:\Unidades compartidas\Informática\Template cartas devolución jubilados\{desdentadoNoNombreDocx}")
                    except Exception:
                        casosInvalidos.append(index)
                        continue

                    

                     # Abrir plantilla y reemplazar
                    #doc = Document("plantilla.docx")
                    docx_replace_regex(doc, re.compile(r"FechaDeEmicionTemplate"), fecha_hoy)
                    docx_replace_regex(doc, re.compile(r"NombreTemplate"), nombre)
                    docx_replace_regex(doc, re.compile(r"RutTemplate"), rut)
                    docx_replace_regex(doc, re.compile(r"EdadTemplate"), str(edad))
                    docx_replace_regex(doc, re.compile(r"FechaDeNacimientoTemplate"), fechaFormateada)
                    docx_replace_regex(doc, re.compile(r"IdTemplate"),str(idTemplate))
                    

                    try:
                        idTemplate+=1
                    except Exception:
                        idInvalidos.append(index)
                        continue


                    # Guardar en carpeta "word"
                    nombre_archivo_seguro = limpiar_nombre_archivo(nombre)
                    ruta_guardado = os.path.join(word_folder, f"{nombre_archivo_seguro} {rut}.docx")
                    doc.save(ruta_guardado)
                    
                    
                   
                    

                except Exception:
                    rutInvalidos.append(index)
                    continue
            else:
                fechasInvalidas.append(index)

    print("\n--- Resumen ---")
    print("Fechas inválidas:", fechasInvalidas)
    print("RUT inválidos:", rutInvalidos)
    print("Edades inválidas:", edadInvalidas)
    print("Fechas válidas:", fechasValidas)
    print("Total pacientes válidos:", val)
    print(f"\nArchivos Word generados en: {os.path.abspath(word_folder)}")
    

else:
    print("No se seleccionó ningún archivo.")




print("\nConvirtiendo archivos Word a PDF...")

#print("WORD FOLDER ORIGINAL", word_folder)
#word_folder = r"C:\PatricioTorres\script_automatizacion_word_pdf\07-05-2025 carta devolución jubilados (2)\07-05-2025 word"
#print("WORD FOLDER FORMATEADO", word_folder)
#pdf_folder = r"C:\PatricioTorres\script_automatizacion_word_pdf\07-05-2025 carta devolución jubilados (2)\07-05-2025 pdf"

# Crear objeto de Word
word = win32com.client.Dispatch("Word.Application")
word.Visible = False

# Convertir todos los .docx a .pdf
for filename in os.listdir(word_folder):
    if filename.endswith(".docx") and not filename.startswith("~$"):  # Evitar archivos temporales
        #docx_path = os.path.join(word_folder, filename)
        docx_path = os.path.abspath(os.path.join(word_folder, filename))
        pdf_name = os.path.splitext(filename)[0] + ".pdf"
        pdf_path = os.path.abspath(os.path.join(pdf_folder, pdf_name))

        try:
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF
            doc.Close()
            print(f"Convertido: {filename}")
        except Exception as e:
            print(f"Error al convertir {filename}: {e}")

# Cerrar Word
word.Quit()


print(f"\n✅ Archivos PDF guardados en: {os.path.abspath(pdf_folder)}")

uc_informatica = rf"G:\Unidades compartidas\Informática\Template cartas devolución jubilados\correlativo.txt"

with open(uc_informatica, "w", encoding="utf-8") as archivo:
    archivo.write(str(idTemplate))

print(f"✅ Archivo creado: {uc_informatica}")



